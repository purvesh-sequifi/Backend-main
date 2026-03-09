#!/usr/bin/env python3
"""
Generate Override Pool System Use Case Review — Word Document
Sequifi | PRD v2.0 | March 2026
"""

import sys
import subprocess
import os

# ── Auto-install python-docx if missing ─────────────────────────────
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENTATION
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENTATION
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

# ════════════════════════════════════════════════════════════════
# DATA
# ════════════════════════════════════════════════════════════════

GROUPS = {
    "FIXED":  {"label": "CALC TYPE 1 — FIXED POOL",   "color": "2E6DA4", "light": "DCE8F6"},
    "GAV":    {"label": "CALC TYPE 2 — % OF GAV",      "color": "2E8B57", "light": "D4EDDA"},
    "XMINUS": {"label": "CALC TYPE 3 — X-MINUS COMM.", "color": "D35400", "light": "FDE9D9"},
    "VOLUME": {"label": "CALC TYPE 4 — VOLUME-TIER",   "color": "6C3483", "light": "E8D5F5"},
}

USE_CASES = [
    # ── GROUP 1: FIXED POOL ──────────────────────────────────────────
    {
        "group": "FIXED", "id": "UC-F1", "calc_type": "Fixed Pool", "dist_mode": "Fixed %",
        "pool_formula": "pool = $600\nL1 = 50% × $600\nL2 = 30% × $600\nL3 = 20% × $600",
        "pool_amount": "$600.00",
        "l1": "$300.00", "l2": "$180.00", "l3": "$120.00",
        "total": "$600.00", "remaining": "$0.00", "rem_status": "ok",
        "efficiency": "100%", "key_risk": "None — full distribution",
        "questions": [
            "Is the $600 fixed value locked permanently per position/product or reviewed periodically? Who has authority to change it?",
            "If a user in the chain is terminated, should their level's percentage be redistributed upward/downward, or simply forfeited?",
            "Should the system warn admin if level percentages don't sum to exactly 100%, leaving an unintentional undistributed amount?",
        ],
    },
    {
        "group": "FIXED", "id": "UC-F2", "calc_type": "Fixed Pool", "dist_mode": "MLM Remaining",
        "pool_formula": "pool = $600\nL1 = 50% × $600\nL2 = 40% × $300\nL3 = 30% × $180",
        "pool_amount": "$600.00",
        "l1": "$300.00", "l2": "$120.00", "l3": "$54.00",
        "total": "$474.00", "remaining": "$126.00", "rem_status": "warn",
        "efficiency": "79%", "key_risk": "$126 remainder fate undefined",
        "questions": [
            "What happens to the $126 undistributed amount? Options: (a) retained by company, (b) rolls over to next period, (c) redistributed to L1 — please confirm business rule.",
            "If the L3 position is vacant (no user configured), does the $54 cascade further down or stop and become part of the remainder?",
            "Should there be a minimum payout threshold (e.g., payout < $1 not written to override record) to avoid tiny bookkeeping entries?",
        ],
    },
    {
        "group": "FIXED", "id": "UC-F3", "calc_type": "Fixed Pool", "dist_mode": "Gap-Fill Diff.",
        "pool_formula": "pool_val = $600/sale\nL1 = 10% × $600\nL2 = 12% gap × $600\nL3 = 13% gap × $600",
        "pool_amount": "$600.00",
        "l1": "$60.00", "l2": "$72.00", "l3": "$78.00",
        "total": "$210.00", "remaining": "$390.00", "rem_status": "high",
        "efficiency": "35%", "key_risk": "65% undistributed — large waste",
        "questions": [
            "$390 (65% of pool) is undistributed with only 3 levels — is this acceptable? Should admin be warned that pool_value ($600) far exceeds what gap-fill will distribute?",
            "Gap-Fill requires tier% per user — who defines these tier percentages when calc type is Fixed (not Volume-Tier)? Is a separate tier table still required?",
            "Should the system block saving a Fixed + Gap-Fill config if no tier table is provided, since distribution depends entirely on tier%?",
        ],
    },
    # ── GROUP 2: % OF GAV ────────────────────────────────────────────
    {
        "group": "GAV", "id": "UC-G1", "calc_type": "% of GAV", "dist_mode": "Fixed %",
        "pool_formula": "pool = $10,000 × 6%\nL1 = 50% × $600\nL2 = 30% × $600\nL3 = 20% × $600",
        "pool_amount": "$600.00",
        "l1": "$300.00", "l2": "$180.00", "l3": "$120.00",
        "total": "$600.00", "remaining": "$0.00", "rem_status": "ok",
        "efficiency": "100%", "key_risk": "Clawback on reversal not defined",
        "questions": [
            "If a sale is reversed/cancelled after commission is paid, should pool overrides be automatically clawed back, or is it a manual adjustment?",
            "Should there be a cap on pool amount (e.g., max $2,000 regardless of GAV) to prevent over-payment on unusually large deals?",
            "Is pool_pct configurable per product within the same position, or is one percentage used across all products for that position?",
        ],
    },
    {
        "group": "GAV", "id": "UC-G2", "calc_type": "% of GAV", "dist_mode": "MLM Remaining",
        "pool_formula": "pool = $10,000 × 6%\nL1 = 50% × $600\nL2 = 40% × $300\nL3 = 30% × $180",
        "pool_amount": "$600.00",
        "l1": "$300.00", "l2": "$120.00", "l3": "$54.00",
        "total": "$474.00", "remaining": "$126.00", "rem_status": "warn",
        "efficiency": "79%", "key_risk": "GAV adjustment handling undefined",
        "questions": [
            "If GAV is adjusted post-calculation (partial chargeback, $10,000 → $7,000), how are already-written override records corrected — reversal + rewrite, or delta record?",
            "Are MLM cascade rates (40% of rem, 30% of rem) configurable per level, or fixed system-wide? Can admin set Level 2 to 60% of remainder?",
            "What is the approved business rule for the $126 remainder? (Same as UC-F2 — needs a consistent answer across all MLM configs)",
        ],
    },
    {
        "group": "GAV", "id": "UC-G3", "calc_type": "% of GAV", "dist_mode": "Gap-Fill Diff.",
        "pool_formula": "pool_val = $10,000 × 6%\nL1 = 10% × $600\nL2 = 12% gap × $600\nL3 = 13% gap × $600",
        "pool_amount": "$600.00",
        "l1": "$60.00", "l2": "$72.00", "l3": "$78.00",
        "total": "$210.00", "remaining": "$390.00", "rem_status": "high",
        "efficiency": "35%", "key_risk": "65% undistributed",
        "questions": [
            "If GAV = $0 (full sale reversal), pool = $0 — should the system create a negative/reversal override record to claw back previously distributed amounts?",
            "What if two users in the chain have the same tier% (gap = 0%) — is a $0 override record created, or is that level silently skipped?",
            "Should pool_pct (6%) apply to adjusted or gross GAV when there are dealer fees, installation costs, or rebates involved?",
        ],
    },
    # ── GROUP 3: X-MINUS ─────────────────────────────────────────────
    {
        "group": "XMINUS", "id": "UC-X1", "calc_type": "X-Minus", "dist_mode": "Fixed %",
        "pool_formula": "pool = max(0, $800-$200)\nL1 = 50% × $600\nL2 = 30% × $600\nL3 = 20% × $600",
        "pool_amount": "$600.00",
        "l1": "$300.00", "l2": "$180.00", "l3": "$120.00",
        "total": "$600.00", "remaining": "$0.00", "rem_status": "ok",
        "efficiency": "100%", "key_risk": "Pool collapses to $0 if rep over-earns",
        "questions": [
            "When pool = $0 because rep over-earned, should this be a reportable event visible in admin reports, or silently skipped with no audit trail?",
            "Is the X value ($800) the same for all reps at this position, or should it be configurable per rep seniority or plan tier?",
            "Should the system alert admin when X is consistently below rep commissions across multiple sales, indicating a possible misconfiguration?",
        ],
    },
    {
        "group": "XMINUS", "id": "UC-X2", "calc_type": "X-Minus", "dist_mode": "MLM Remaining",
        "pool_formula": "pool = max(0, $800-$200)\nL1 = 50% × $600\nL2 = 40% × $300\nL3 = 30% × $180",
        "pool_amount": "$600.00",
        "l1": "$300.00", "l2": "$120.00", "l3": "$54.00",
        "total": "$474.00", "remaining": "$126.00", "rem_status": "warn",
        "efficiency": "79%", "key_risk": "Entire cascade = $0 when rep over-earns",
        "questions": [
            "When pool = $0 (rep over-earns), the entire MLM cascade produces nothing — should admins receive a proactive alert during payroll review, not after?",
            "Can the X value be dynamic (e.g., X = GAV × 10%) rather than a fixed dollar amount, to scale with deal size?",
            "For quarterly payout triggers, is X-Minus applied per individual sale or against total quarter commission vs. total quarter X budget?",
        ],
    },
    {
        "group": "XMINUS", "id": "UC-X3", "calc_type": "X-Minus", "dist_mode": "Gap-Fill Diff.",
        "pool_formula": "pool_val = max(0,$800-$200)\nL1 = 10% × $600\nL2 = 12% gap × $600\nL3 = 13% gap × $600",
        "pool_amount": "$600.00",
        "l1": "$60.00", "l2": "$72.00", "l3": "$78.00",
        "total": "$210.00", "remaining": "$390.00", "rem_status": "high",
        "efficiency": "35%", "key_risk": "Double risk: X-Minus + 65% undistributed",
        "questions": [
            "Gap-Fill already leaves 65% undistributed ($390). If X-Minus further reduces the pool, should the simulation tool prominently flag the compounding impact before admin commits?",
            "Is it valid business logic to combine X-Minus (commission-driven) with Gap-Fill (tier-driven)? Should this combination require explicit admin confirmation?",
            "Should there be a minimum gap-fill payout floor (e.g., $5/level) to avoid writing $3–$4 override records that add bookkeeping noise?",
        ],
    },
    # ── GROUP 4: VOLUME-TIER ─────────────────────────────────────────
    {
        "group": "VOLUME", "id": "UC-V1", "calc_type": "Volume-Tier", "dist_mode": "Fixed %",
        "pool_formula": "pool = 300×$50 = $15,000\nL1 = 50% × $15,000\nL2 = 30% × $15,000\nL3 = 20% × $15,000",
        "pool_amount": "$15,000.00",
        "l1": "$7,500.00", "l2": "$4,500.00", "l3": "$3,000.00",
        "total": "$15,000.00", "remaining": "$0.00", "rem_status": "ok",
        "efficiency": "100%", "key_risk": "Tier% ignored in this mode",
        "questions": [
            "In this mode, individual tier% (10%, 22%, 35%) is not used for distribution — the fixed level% overrides it. Should the UI explicitly warn admin that tier percentages are ignored?",
            "If a new rep joins mid-quarter with 0 downline, do they affect their upline's downline count (tier%) immediately, or is the snapshot taken at quarter-end?",
            "Should the quarterly batch job be idempotent (safe to re-run without duplicating overrides)? If run twice, will it double-write $15,000?",
        ],
    },
    {
        "group": "VOLUME", "id": "UC-V2", "calc_type": "Volume-Tier", "dist_mode": "MLM Remaining",
        "pool_formula": "pool = 300×$50 = $15,000\nL1 = 50% × $15,000\nL2 = 40% × $7,500\nL3 = 30% × $4,500",
        "pool_amount": "$15,000.00",
        "l1": "$7,500.00", "l2": "$3,000.00", "l3": "$1,350.00",
        "total": "$11,850.00", "remaining": "$3,150.00", "rem_status": "high",
        "efficiency": "79%", "key_risk": "$3,150 remainder at scale is material",
        "questions": [
            "$3,150 undistributed (21% of $15,000) — at quarterly scale this is significant. What is the approved business rule for this remainder? Must be confirmed before go-live.",
            "If L1 is terminated when the quarterly batch runs, does their $7,500 share cascade down to L2, revert to company, or require manual handling?",
            "Can MLM cascade percentages (50% / 40% / 30%) be customized per pool configuration, or are they fixed globally in the system?",
        ],
    },
    {
        "group": "VOLUME", "id": "UC-V3 ★", "calc_type": "Volume-Tier\n(Grow Mktg)", "dist_mode": "Gap-Fill Diff.",
        "pool_formula": "pool = 300×$50 = $15,000\nL1:300×$50×10% = $1,500\nL2:300×$50×12% = $1,800\nL3:300×$50×13% = $1,950",
        "pool_amount": "$15,000.00",
        "l1": "$1,500.00", "l2": "$1,800.00", "l3": "$1,950.00",
        "total": "$5,250.00", "remaining": "$9,750.00", "rem_status": "high",
        "efficiency": "35%", "key_risk": "$9,750 undistributed — confirm with Grow Mktg",
        "questions": [
            "CRITICAL: $9,750 (65% of $15,000) remains undistributed with 3 levels. Is Grow Marketing aware that only 35% of the total pool pot is distributed? Must confirm this matches their manual spreadsheet.",
            "Are tier percentages calculated at quarter-end (using final downline counts) or locked at quarter-start? If a user crosses tier boundary mid-quarter (499→500 sales), do they get higher rate for full quarter or only from crossing date?",
            "If Grow Marketing adds an L4/L5 user to the chain in future, will the system automatically pick up the additional tier differential, or does max_depth need manual updating?",
        ],
    },
]

CRITICAL_ITEMS = [
    ("1", "Undistributed pool remainder rule",
     "UC-F2, G2, X2, X3, V2, V3, F3, G3",
     "Revenue leakage or over-payment depending on rule chosen",
     "Retain by company / roll over to next period / redistribute to L1?"),
    ("2", "Grow Marketing 65% undistributed (UC-V3 only distributes $5,250 of $15,000)",
     "UC-V3 — Grow Marketing",
     "Customer validation may fail parallel run — blocks go-live",
     "Confirm with Grow Marketing that 35% distribution is their intended design"),
    ("3", "Tier% source for non-Volume-Tier gap-fill configs (Fixed, GAV, X-Minus + Gap-Fill)",
     "UC-F3, G3, X3",
     "System cannot compute gap-fill without tier% — will error at runtime",
     "Must a tier table always be configured when Gap-Fill mode is selected?"),
    ("4", "Sale reversal / clawback handling for % of GAV configs",
     "UC-G1, G2, G3",
     "Overrides remain on record after sale reversal — incorrect payroll output",
     "Auto-reverse override records on GAV adjustment, or manual process?"),
    ("5", "Terminated user in chain — forfeited level percentage",
     "All 12 UCs",
     "Chain members above terminated user receive less than configured",
     "Redistribute forfeited % upward, or is forfeit the intended behavior?"),
    ("6", "Quarterly batch idempotency — double-run risk for Volume-Tier",
     "UC-V1, V2, V3",
     "Double-writes $15,000 in overrides if batch triggered twice",
     "Add unique constraint on (sale_id, user_id, period, overrides_type='pool')?"),
]

# ════════════════════════════════════════════════════════════════
# HELPERS
# ════════════════════════════════════════════════════════════════

def rgb(hex_str):
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def set_col_width(table, col_idx, width_inches):
    for row in table.rows:
        row.cells[col_idx].width = Inches(width_inches)


def cell_para(cell, text, size=8.5, bold=False, italic=False,
              color="000000", align=WD_ALIGN_PARAGRAPH.LEFT, first=True):
    if first:
        para = cell.paragraphs[0]
    else:
        para = cell.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = rgb(color)
    return para


def add_header_row(table, headers, bg="1B3A6B", text_color="FFFFFF", font_size=8):
    row = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = row.cells[i]
        shade_cell(cell, bg)
        cell_para(cell, hdr, size=font_size, bold=True, color=text_color,
                  align=WD_ALIGN_PARAGRAPH.CENTER)


def border_table(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "B0B8CC")
        borders.append(b)
    tblPr.append(borders)


# ════════════════════════════════════════════════════════════════
# DOCUMENT BUILD
# ════════════════════════════════════════════════════════════════

def build_document():
    doc = Document()

    # ── Page setup: Landscape A4 ────────────────────────────────
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width  = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin   = Cm(1.2)
    section.right_margin  = Cm(1.2)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # ── Title ────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("SEQUIFI — Override Pool System")
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = rgb("1B3A6B")

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(
        "Use Case Review Sheet  ·  All 12 Calculation × Distribution Combinations  ·  PRD v2.0"
    )
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = rgb("4A5568")

    # ── Meta table ───────────────────────────────────────────────
    meta = doc.add_table(rows=1, cols=5)
    meta.style = "Table Grid"
    border_table(meta)
    meta_headers = ["Document Version", "PRD Reference", "Prepared By", "Date", "Status"]
    meta_values  = ["1.0", "Override Pool PRD v2.0", "Engineering Team", "March 2026", "⚠  Awaiting Product Owner Review"]
    hrow = meta.rows[0]
    for i, (h, v) in enumerate(zip(meta_headers, meta_values)):
        c = hrow.cells[i]
        shade_cell(c, "1B3A6B")
        cell_para(c, h, size=8, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
    vrow = meta.add_row()
    for i, v in enumerate(meta_values):
        c = vrow.cells[i]
        shade_cell(c, "EEF2F8")
        col = "C0392B" if i == 4 else "1A1A2E"
        cell_para(c, v, size=8.5, bold=(i == 4), color=col, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # ── Parameters table ─────────────────────────────────────────
    ph = doc.add_paragraph()
    ph.add_run("Base Parameters (Shared Across All 12 Combinations)").font.size = Pt(10)
    ph.runs[0].font.bold = True
    ph.runs[0].font.color.rgb = rgb("1B3A6B")

    ptbl = doc.add_table(rows=2, cols=4)
    ptbl.style = "Table Grid"
    border_table(ptbl)
    param_headers = ["Sale Parameters", "Pool Config", "Fixed % Distribution Rates", "Gap-Fill Tier % (per user downline)"]
    param_vals = [
        "GAV = $10,000\nRep Commission = $200\nX value (X-Minus) = $800",
        "Fixed pool_value = $600\n% of GAV pool_pct = 6%\nVol-Tier pool/sale = $50\nQuarterly sales = 300",
        "Level 1 = 50%\nLevel 2 = 30%\nLevel 3 = 20%\n\nMLM: L1 takes 50%, L2 takes 40% of rem, L3 takes 30% of rem",
        "L1 downline: 300 sales → 10%\nL2 downline: 800 sales → 22%\nL3 downline: 1,500 sales → 35%\n\nChain: Seller → L1 → L2 → L3",
    ]
    for i, h in enumerate(param_headers):
        shade_cell(ptbl.rows[0].cells[i], "2C4F8A")
        cell_para(ptbl.rows[0].cells[i], h, size=8.5, bold=True, color="FFFFFF",
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, v in enumerate(param_vals):
        shade_cell(ptbl.rows[1].cells[i], "F4F7FB")
        cell_para(ptbl.rows[1].cells[i], v, size=8.5)

    doc.add_paragraph()

    # ── Main Use Case table ───────────────────────────────────────
    mh = doc.add_paragraph()
    mh.add_run("Use Case Combinations — Pool Calculation Amounts & Product Owner Questions").font.size = Pt(10)
    mh.runs[0].font.bold = True
    mh.runs[0].font.color.rgb = rgb("1B3A6B")

    # Column headers + widths (landscape ~27.3cm usable)
    COL_HEADERS = [
        "UC ID", "Calc Type", "Dist. Mode", "Pool Formula",
        "Pool Amt", "L1 Amount", "L2 Amount", "L3 Amount",
        "Total Paid", "Remaining", "❓ Questions for Product Owner"
    ]
    COL_WIDTHS = [0.55, 0.75, 0.85, 1.55, 0.65, 0.65, 0.65, 0.65, 0.70, 0.70, 2.85]

    # Count rows: 1 header + 4 group headers + 12 UC rows + 1 edge case = 18
    total_rows = 1 + len(USE_CASES) + 4
    mtbl = doc.add_table(rows=total_rows, cols=11)
    mtbl.style = "Table Grid"
    border_table(mtbl)

    # Header row
    add_header_row(mtbl, COL_HEADERS, bg="1B3A6B", text_color="FFFFFF", font_size=8)
    for ci, w in enumerate(COL_WIDTHS):
        mtbl.rows[0].cells[ci].width = Inches(w)

    row_idx = 1
    prev_group = None

    for uc in USE_CASES:
        grp = uc["group"]
        ginfo = GROUPS[grp]

        # ── Group header row ──────────────────────────────────────
        if grp != prev_group:
            grow = mtbl.rows[row_idx]
            merged = grow.cells[0].merge(grow.cells[10])
            shade_cell(merged, ginfo["light"])
            cell_para(
                merged,
                f"  {ginfo['label']}   —   Pool formula varies per calc type; distribution mode determines how that pool is split across L1→L2→L3",
                size=8.5, bold=True, color=ginfo["color"]
            )
            row_idx += 1
            prev_group = grp

        # ── UC data row ───────────────────────────────────────────
        urow = mtbl.rows[row_idx]
        for ci, w in enumerate(COL_WIDTHS):
            urow.cells[ci].width = Inches(w)

        # UC ID cell
        shade_cell(urow.cells[0], ginfo["color"])
        cell_para(urow.cells[0], uc["id"], size=8, bold=True, color="FFFFFF",
                  align=WD_ALIGN_PARAGRAPH.CENTER)

        # Calc Type
        shade_cell(urow.cells[1], ginfo["light"])
        cell_para(urow.cells[1], uc["calc_type"], size=8, bold=True, color=ginfo["color"])

        # Dist Mode badge
        dist_colors = {
            "Fixed %":        ("E3F0FF", "1B3A6B"),
            "MLM Remaining":  ("FFF0E3", "B84000"),
            "Gap-Fill Diff.": ("F0E3FF", "5B1A8A"),
        }
        dbg, dfc = dist_colors.get(uc["dist_mode"], ("F5F5F5", "333333"))
        shade_cell(urow.cells[2], dbg)
        cell_para(urow.cells[2], uc["dist_mode"], size=8, bold=True, color=dfc,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

        # Pool formula (monospace feel via italic)
        shade_cell(urow.cells[3], "F8FAFE")
        cell_para(urow.cells[3], uc["pool_formula"], size=8, italic=True, color="2D3A52")

        # Pool amount
        shade_cell(urow.cells[4], "EEF2F8")
        cell_para(urow.cells[4], uc["pool_amount"], size=8.5, bold=True, color="1B3A6B",
                  align=WD_ALIGN_PARAGRAPH.RIGHT)

        # L1 L2 L3
        for ci_off, key in enumerate(["l1", "l2", "l3"], start=5):
            shade_cell(urow.cells[ci_off], "FAFBFF")
            cell_para(urow.cells[ci_off], uc[key], size=8.5, bold=True, color="2D3A52",
                      align=WD_ALIGN_PARAGRAPH.RIGHT)

        # Total
        shade_cell(urow.cells[8], "E8F0FE")
        cell_para(urow.cells[8], uc["total"], size=8.5, bold=True, color="1B3A6B",
                  align=WD_ALIGN_PARAGRAPH.RIGHT)

        # Remaining
        rem_cfg = {
            "ok":   ("EAFFEA", "27AE60"),
            "warn": ("FFF8E6", "E67E22"),
            "high": ("FFF0F0", "C0392B"),
        }
        rbg, rfc = rem_cfg[uc["rem_status"]]
        shade_cell(urow.cells[9], rbg)
        suffix = " ✓" if uc["rem_status"] == "ok" else (" ⚠" if uc["rem_status"] == "warn" else " !")
        cell_para(urow.cells[9], uc["remaining"] + suffix, size=8.5, bold=True, color=rfc,
                  align=WD_ALIGN_PARAGRAPH.RIGHT)

        # Questions
        shade_cell(urow.cells[10], "FFFEF0")
        for qi, q in enumerate(uc["questions"]):
            is_first = qi == 0
            bold_q = q.startswith("CRITICAL")
            color_q = "C0392B" if bold_q else "3D2C00"
            prefix = f"Q{qi+1}. "
            cell_para(urow.cells[10], prefix + q, size=8, bold=bold_q,
                      color=color_q, first=is_first)

        row_idx += 1

    # Edge case note row
    erow = mtbl.rows[row_idx] if row_idx < total_rows else None
    if erow:
        merged_e = erow.cells[0].merge(erow.cells[10])
        shade_cell(merged_e, "FFF3F3")
        cell_para(
            merged_e,
            "⚠  EDGE CASE — X-Minus only:  If rep commission ($900) > X value ($800)"
            "  →  pool = max(0, $800−$900) = $0  →  No override records written at any level (L1, L2, L3 all receive $0)."
            "  Q: Should this be logged in the batch audit log so admin can investigate why recruiters received nothing for a given sale?",
            size=8, italic=True, color="7B1A1A"
        )

    doc.add_paragraph()

    # ── Summary table ─────────────────────────────────────────────
    sh = doc.add_paragraph()
    sh.add_run("Summary — Distribution Efficiency Across All 12 Combinations").font.size = Pt(10)
    sh.runs[0].font.bold = True
    sh.runs[0].font.color.rgb = rgb("1B3A6B")

    sum_headers = ["UC ID", "Calc Type", "Distribution Mode", "Pool Amt", "L1", "L2", "L3", "Total Paid", "Remaining", "Efficiency", "Key Risk"]
    sum_widths  = [0.55, 0.75, 0.9, 0.75, 0.65, 0.65, 0.65, 0.75, 0.75, 0.65, 2.5]

    stbl = doc.add_table(rows=1 + len(USE_CASES), cols=11)
    stbl.style = "Table Grid"
    border_table(stbl)
    add_header_row(stbl, sum_headers, bg="1B3A6B", text_color="FFFFFF", font_size=8)
    for ci, w in enumerate(sum_widths):
        stbl.rows[0].cells[ci].width = Inches(w)

    for ri, uc in enumerate(USE_CASES, start=1):
        row = stbl.rows[ri]
        for ci, w in enumerate(sum_widths):
            row.cells[ci].width = Inches(w)

        eff_val = uc["efficiency"]
        is_best  = eff_val == "100%"
        is_worst = uc["rem_status"] == "high"
        row_bg = "EAFFEA" if is_best else ("FFF5F5" if is_worst else "FAFAFA")

        data = [
            uc["id"], uc["calc_type"].split("\n")[0], uc["dist_mode"],
            uc["pool_amount"], uc["l1"], uc["l2"], uc["l3"],
            uc["total"], uc["remaining"], eff_val, uc["key_risk"]
        ]
        for ci, val in enumerate(data):
            shade_cell(row.cells[ci], row_bg)
            align = WD_ALIGN_PARAGRAPH.RIGHT if ci in (3,4,5,6,7,8) else WD_ALIGN_PARAGRAPH.CENTER if ci in (0,9) else WD_ALIGN_PARAGRAPH.LEFT
            bold_cell = ci in (7, 9)
            color_cell = "27AE60" if (ci == 9 and eff_val == "100%") else ("C0392B" if (ci == 9 and eff_val == "35%") else "E67E22" if ci == 9 else "1A1A2E")
            cell_para(row.cells[ci], val, size=8, bold=bold_cell, color=color_cell, align=align)

    doc.add_paragraph()

    # ── Critical Open Items ───────────────────────────────────────
    ch = doc.add_paragraph()
    ch.add_run("🚨  Critical Open Items — Product Owner Decision Required Before Implementation").font.size = Pt(10)
    ch.runs[0].font.bold = True
    ch.runs[0].font.color.rgb = rgb("C0392B")

    ci_headers = ["#", "Issue", "Affects UCs", "Impact if Unresolved", "Decision Required"]
    ci_widths  = [0.25, 2.8, 1.2, 2.2, 2.45]

    ctbl = doc.add_table(rows=1 + len(CRITICAL_ITEMS), cols=5)
    ctbl.style = "Table Grid"
    border_table(ctbl)
    add_header_row(ctbl, ci_headers, bg="C0392B", text_color="FFFFFF", font_size=8)
    for ci, w in enumerate(ci_widths):
        ctbl.rows[0].cells[ci].width = Inches(w)

    for ri, item in enumerate(CRITICAL_ITEMS, start=1):
        row = ctbl.rows[ri]
        for ci, w in enumerate(ci_widths):
            row.cells[ci].width = Inches(w)
        row_bg = "FFF5F5" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(item):
            shade_cell(row.cells[ci], row_bg)
            align = WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT
            cell_para(row.cells[ci], val, size=8, bold=(ci == 0), color="1A1A2E", align=align)

    # ── Footer note ───────────────────────────────────────────────
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_para.add_run(
        "Sequifi — Override Pool System  ·  Use Case Review v1.0  ·  "
        "PRD Reference: Override Pool PRD v2.0 (Feb 2026)  ·  "
        "Total Combinations: 12  ·  Open Questions: 36  ·  Critical Items: 6  ·  "
        "Please annotate responses and return to Engineering team"
    )
    fr.font.size = Pt(7.5)
    fr.font.italic = True
    fr.font.color.rgb = rgb("6B7A99")

    # ── Save ──────────────────────────────────────────────────────
    out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                            "override_pool_usecase_review.docx")
    doc.save(out_path)
    print(f"✅  Document saved: {out_path}")
    return out_path


if __name__ == "__main__":
    build_document()
